# utils/cv_parser.py
"""
ATS CV Parser Utility
Ekstrak data dari file CV (PDF / DOCX) secara otomatis.

Dependensi (install salah satu untuk PDF):
    pip install PyMuPDF          # Rekomendasi
    pip install pdfplumber
    pip install pypdf

Wajib untuk DOCX:
    pip install python-docx
"""

import re


# ─────────────────────────────────────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(filepath: str) -> str:
    """Ekstrak teks mentah dari file PDF."""
    # Coba PyMuPDF (fitz) — paling akurat
    try:
        import fitz
        doc = fitz.open(filepath)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    except ImportError:
        pass

    # Fallback pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            return "\n".join(
                p.extract_text() for p in pdf.pages if p.extract_text()
            )
    except ImportError:
        pass

    # Fallback pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(filepath)
        return "\n".join(
            p.extract_text() for p in reader.pages if p.extract_text()
        )
    except ImportError:
        pass

    raise ImportError(
        "Tidak ada library PDF yang tersedia.\n"
        "Install salah satu: pip install PyMuPDF  /  pip install pdfplumber"
    )


def extract_text_from_docx(filepath: str) -> str:
    """Ekstrak teks mentah dari file DOCX."""
    try:
        from docx import Document
        doc = Document(filepath)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)
    except ImportError:
        raise ImportError(
            "Library python-docx tidak tersedia.\n"
            "Install dengan: pip install python-docx"
        )


def extract_text(filepath: str) -> str:
    """Auto-detect format dan ekstrak teks."""
    lower = filepath.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(filepath)
    elif lower.endswith(".docx"):
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(
            f"Format tidak didukung: {filepath}\n"
            "Hanya PDF (.pdf) dan Word (.docx) yang diterima."
        )


# ─────────────────────────────────────────────────────────────────────────────
# PARSER
# ─────────────────────────────────────────────────────────────────────────────

class CVParser:
    """
    Parse teks CV dan kembalikan dict terstruktur:
    {
        nama, email, no_hp, pendidikan (str level tertinggi),
        pengalaman_tahun (int), skill (str),
        pendidikan_detail (list), pengalaman_detail (list),
    }
    """

    SECTION_EDUCATION = [
        'pendidikan', 'education', 'riwayat pendidikan',
        'academic background', 'educational background',
    ]
    SECTION_EXPERIENCE = [
        'pengalaman', 'experience', 'work experience',
        'pengalaman kerja', 'riwayat pekerjaan', 'employment history',
    ]
    SECTION_SKILLS = [
        'keahlian', 'skill', 'skills', 'kemampuan',
        'kompetensi', 'technical skills', 'keterampilan',
    ]

    # Mapping keyword → level pendidikan Django (pilihan di model)
    EDU_LEVEL_MAP = [
        (['s3', 'doktor', 'phd', 'doctor'],                         'S3'),
        (['s2', 'magister', 'master', 'mba'],                       'S2'),
        (['s1', 'sarjana', 'bachelor', 'strata 1', 'strata-1'],     'S1'),
        (['d4'],                                                      'S1'),
        (['d3', 'diploma 3', 'diploma iii'],                         'D3'),
        (['d2', 'd1'],                                               'D3'),
        (['smk', 'sma', 'ma ', 'aliyah', 'high school', 'vocational'], 'SMA/SMK'),
        (['smp', 'mts'],                                             'SMP'),
        (['sd', 'sekolah dasar'],                                    'SD'),
    ]

    def parse(self, text: str) -> dict:
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        sections = self._split_sections(lines)

        edu_detail  = self._parse_education(sections.get('education', []))
        exp_detail  = self._parse_experience(sections.get('experience', []))
        skill_list  = self._parse_skills(sections.get('skills', []))

        return {
            # Field langsung ke model Candidate
            'nama':              self._extract_name(lines),
            'email':             self._extract_email(text),
            'no_hp':             self._extract_phone(text),
            'pendidikan':        self._highest_education(edu_detail),
            'pengalaman_tahun':  self._total_experience_years(exp_detail),
            'skill':             ', '.join(skill_list[:20]),  # simpan sebagai string

            # Detail untuk ditampilkan di UI (tidak langsung ke model)
            'pendidikan_detail':  edu_detail,
            'pengalaman_detail':  exp_detail,
            'skill_list':         skill_list,
        }

    # ── DATA PRIBADI ─────────────────────────────────────────────────────────

    def _extract_name(self, lines: list) -> str:
        skip = ['curriculum', 'vitae', 'resume', 'cv', 'profile',
                'data diri', 'biodata', 'personal', 'profil', 'daftar riwayat']
        for line in lines[:8]:
            low = line.lower()
            if any(k in low for k in skip):
                continue
            if re.search(r'[@:/\\]|http|\.com|\.id|\d{5,}', low):
                continue
            words = line.split()
            if 2 <= len(words) <= 6 and all(re.match(r"[A-Za-z.,'\-]+$", w) for w in words):
                return line.title()
        return ''

    def _extract_email(self, text: str) -> str:
        m = re.search(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}', text)
        return m.group(0) if m else ''

    def _extract_phone(self, text: str) -> str:
        # +62 / 08 / 62
        m = re.search(r'(\+?62[\s\-]?|0)[0-9][\d\s\-]{7,13}', text)
        if m:
            return re.sub(r'[\s\-]', '', m.group(0))
        m2 = re.search(r'\b0\d{9,12}\b', text)
        return m2.group(0) if m2 else ''

    # ── SECTIONS ─────────────────────────────────────────────────────────────

    def _split_sections(self, lines: list) -> dict:
        sections = {'education': [], 'experience': [], 'skills': [], 'other': []}
        current = 'other'
        for line in lines:
            low = line.lower().strip(':').strip()
            if any(k == low or k in low for k in self.SECTION_EDUCATION):
                current = 'education'; continue
            elif any(k == low or k in low for k in self.SECTION_EXPERIENCE):
                current = 'experience'; continue
            elif any(k == low or k in low for k in self.SECTION_SKILLS):
                current = 'skills'; continue
            sections[current].append(line)
        return sections

    # ── PENDIDIKAN ───────────────────────────────────────────────────────────

    def _parse_education(self, lines: list) -> list:
        entries, current = [], {}
        degree_kw = ['s1','s2','s3','d1','d2','d3','d4','sarjana','magister',
                     'doktor','diploma','bachelor','master','phd','sma','smk',
                     'smp','sd','high school','vocational']
        inst_kw   = ['universitas','university','institut','sekolah','college',
                     'politeknik','academy','stie','stmik','uin','iain']

        for line in lines:
            low = line.lower()
            yr  = re.search(r'(19|20)\d{2}', line)
            is_deg  = any(k in low for k in degree_kw)
            is_inst = any(k in low for k in inst_kw)

            if is_deg or is_inst:
                if current: entries.append(current)
                current = {'institusi': '', 'jurusan': '', 'jenjang': '', 'tahun': ''}

            if current:
                if is_inst and not current['institusi']:
                    current['institusi'] = line
                elif is_deg and not current['jenjang']:
                    current['jenjang'] = line
                elif yr and not current['tahun']:
                    current['tahun'] = yr.group(0)
                elif not current['jurusan'] and not is_inst and not is_deg:
                    current['jurusan'] = line

        if current: entries.append(current)
        return entries

    def _highest_education(self, edu_list: list) -> str:
        """Kembalikan level pendidikan tertinggi (sesuai choices di model)."""
        found_levels = []
        order = ['S3','S2','S1','D3','SMA/SMK','SMP','SD']
        for edu in edu_list:
            text = (edu.get('jenjang','') + ' ' + edu.get('institusi','')).lower()
            for keywords, level in self.EDU_LEVEL_MAP:
                if any(k in text for k in keywords):
                    found_levels.append(level)
                    break
        for level in order:
            if level in found_levels:
                return level
        return ''

    # ── PENGALAMAN ───────────────────────────────────────────────────────────

    def _parse_experience(self, lines: list) -> list:
        entries, current = [], {}
        company_kw  = ['pt.','pt ','cv.','cv ','tbk','ltd','inc','corp','co.',
                       'perusahaan','company','group']
        position_kw = ['staff','manager','supervisor','direktur','kepala',
                       'koordinator','analyst','engineer','developer','programmer',
                       'designer','officer','admin','specialist','consultant',
                       'intern','magang','junior','senior','lead','head']

        for line in lines:
            low = line.lower()
            yr  = re.search(
                r'(20\d{2}|19\d{2})\s*[-–]\s*(20\d{2}|19\d{2}|sekarang|present|now)',
                line, re.I)
            is_company  = any(k in low for k in company_kw)
            is_position = any(k in low for k in position_kw)

            if is_company:
                if current: entries.append(current)
                current = {'perusahaan': line, 'posisi': '', 'durasi': '', 'tahun_mulai': 0, 'tahun_selesai': 0}
            elif is_position and current and not current['posisi']:
                current['posisi'] = line
            elif yr:
                if not current:
                    current = {'perusahaan': '', 'posisi': '', 'durasi': '', 'tahun_mulai': 0, 'tahun_selesai': 0}
                current['durasi'] = yr.group(0)
                # Hitung tahun untuk kalkulasi total pengalaman
                start_m = re.search(r'(19|20)\d{2}', yr.group(0))
                end_raw = re.search(r'[-–]\s*(.+)$', yr.group(0), re.I)
                if start_m:
                    current['tahun_mulai'] = int(start_m.group(0))
                if end_raw:
                    end_str = end_raw.group(1).strip().lower()
                    if any(k in end_str for k in ['sekarang','present','now']):
                        import datetime
                        current['tahun_selesai'] = datetime.date.today().year
                    else:
                        end_m = re.search(r'(19|20)\d{2}', end_str)
                        if end_m:
                            current['tahun_selesai'] = int(end_m.group(0))

        if current: entries.append(current)
        return entries

    def _total_experience_years(self, exp_list: list) -> int:
        """Hitung total pengalaman kerja dalam tahun."""
        total = 0
        for exp in exp_list:
            start = exp.get('tahun_mulai', 0)
            end   = exp.get('tahun_selesai', 0)
            if start and end and end >= start:
                total += (end - start)
        return min(total, 50)  # sanity cap

    # ── SKILL ─────────────────────────────────────────────────────────────────

    def _parse_skills(self, lines: list) -> list:
        full = ' '.join(lines)
        parts = re.split(r'[,|•·\-\n/]', full)
        seen, skills = set(), []
        for p in parts:
            s = p.strip()
            if 2 <= len(s) <= 50 and s.lower() not in seen:
                seen.add(s.lower())
                skills.append(s)
        return skills